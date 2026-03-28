from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_styled(doc, text, level):
    para = doc.add_heading(text, level=level)
    return para

def build_ddr_report(ddr_data, inspection_images, thermal_images, output_path):
    doc = Document()
    
    # Page margins set karo
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Property info
    doc.add_paragraph("Inspection Date: 27.09.2022")
    doc.add_paragraph("Inspected By: Krushna & Mahesh")
    doc.add_paragraph("Property Type: Flat | Floors: 11")
    doc.add_paragraph("Company: UrbanRoof")
    doc.add_paragraph("")
    
    # ─── Section 1: Property Issue Summary ───
    doc.add_heading("1. Property Issue Summary", level=1)
    doc.add_paragraph(ddr_data.get("property_issue_summary", "Not Available"))
    doc.add_paragraph("")
    
    # ─── Section 2: Area-wise Observations ───
    doc.add_heading("2. Area-wise Observations", level=1)
    areas = ddr_data.get("area_wise_observations", [])
    
    # Inspection images aur thermal images alag alag list mein
    insp_imgs = [img for img in inspection_images if os.path.exists(img["path"])]
    therm_imgs = [img for img in thermal_images if os.path.exists(img["path"])]
    
    # Sirf odd pages thermal images hain (1,3,5...) even pages normal photos hain
    actual_thermal = therm_imgs
    actual_insp = insp_imgs
    
    for i, area in enumerate(areas):
        doc.add_heading(f"Area {i+1}: {area.get('area', 'Unknown')}", level=2)
        
        obs = area.get("observation", "Not Available")
        doc.add_paragraph(f"Observation: {obs}")
        
        thermal = area.get("thermal_finding", "Not Available")
        doc.add_paragraph(f"Thermal Finding: {thermal}")
        
        # Inspection image lagao
        doc.add_paragraph("Inspection Photo:")
        if i < len(insp_imgs):
            try:
                doc.add_picture(insp_imgs[i]["path"], width=Inches(2.5))
                doc.add_paragraph(f"(Source: Inspection Report - Page {insp_imgs[i]['page']})")
            except Exception as e:
                doc.add_paragraph("Image: Not Available")
        else:
            doc.add_paragraph("Image: Not Available")
        
        # Thermal image lagao
        doc.add_paragraph("Thermal Image:")
        if i < len(actual_thermal):
            try:
                doc.add_picture(actual_thermal[i]["path"], width=Inches(2.5))
                doc.add_paragraph(f"(Source: Thermal Report - Page {actual_thermal[i]['page']})")
            except Exception as e:
                doc.add_paragraph("Thermal Image: Not Available")
        else:
            doc.add_paragraph("Thermal Image: Not Available")
        
        doc.add_paragraph("")
    
    # ─── Section 3: Probable Root Cause ───
    doc.add_heading("3. Probable Root Cause", level=1)
    doc.add_paragraph(ddr_data.get("probable_root_cause", "Not Available"))
    doc.add_paragraph("")
    
    # ─── Section 4: Severity Assessment ───
    doc.add_heading("4. Severity Assessment", level=1)
    severity = ddr_data.get("severity_assessment", {})
    level = severity.get("level", "Not Available")
    reasoning = severity.get("reasoning", "Not Available")
    doc.add_paragraph(f"Severity Level: {level}")
    doc.add_paragraph(f"Reasoning: {reasoning}")
    doc.add_paragraph("")
    
    # ─── Section 5: Recommended Actions ───
    doc.add_heading("5. Recommended Actions", level=1)
    actions = ddr_data.get("recommended_actions", [])
    if actions:
        for action in actions:
            doc.add_paragraph(f"• {action}")
    else:
        doc.add_paragraph("Not Available")
    doc.add_paragraph("")
    
    # ─── Section 6: Additional Notes ───
    doc.add_heading("6. Additional Notes", level=1)
    doc.add_paragraph(ddr_data.get("additional_notes", "Not Available"))
    doc.add_paragraph("")
    
    # ─── Section 7: Missing or Unclear Information ───
    doc.add_heading("7. Missing or Unclear Information", level=1)
    missing = ddr_data.get("missing_or_unclear_information", "Not Available")
    doc.add_paragraph(missing)
    doc.add_paragraph("")
    

    
    # Save
    doc.save(output_path)
    print(f"Report saved: {output_path}")
    return output_path